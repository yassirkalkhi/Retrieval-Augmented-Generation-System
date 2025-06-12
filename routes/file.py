from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from services.chroma import ChromaService
from models.schemas import Document, QueryRequest
from pydantic import BaseModel
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PIL import Image
import docx
import uuid
import tempfile
import os
import asyncio
import logging

logger = logging.getLogger(__name__)
router = APIRouter()

db_pdf = ChromaService(collection_name="docs", embedding_type='text')
db_image = ChromaService(collection_name="images", embedding_type="image")
db_docx = ChromaService(collection_name="docs", embedding_type='text')


async def store_pdf(file: UploadFile, uid: str, source: str):
    logger.info(f"Starting PDF processing for file: {file.filename}, UID: {uid}, Source: {source}")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(await file.read())
        path = tmp.name
    logger.info(f"PDF file '{file.filename}' saved to temporary path: {path}")
    
    loader = PyPDFLoader(path)
    pages = [page async for page in loader.alazy_load()]
    logger.info(f"Loaded {len(pages)} pages from '{file.filename}'")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=20,
        length_function=len,
        is_separator_regex=False,
        separators=["\n\n", "\n", ".", ",", "\u200b", "\uff0c", "\u3001", "\uff0e", "\u3002", ""]
    )
    chunks = []

    for i, page in enumerate(pages, start=1):
        metadata = {"page": i, "source": source, "uid": uid}
        split_docs = splitter.create_documents([page.page_content])
        
        for j, chunk in enumerate(split_docs, start=1):
            document = Document(
                content=chunk.page_content,
                metadata={**metadata, "id": str(uuid.uuid4()), "chunk": str(j)}
            )
            chunks.append(document)
    
    if chunks:
        logger.info(f"Storing {len(chunks)} chunks for source: '{source}' in ChromaDB.")
        db_pdf.store(chunks)
        logger.info(f"Successfully stored chunks for source: '{source}'.")
    os.remove(path)
    logger.info(f"Finished PDF processing and removed temporary file for: {file.filename}")

async def store_image(file: UploadFile, uid: str, source: str):
    logger.info(f"Starting image processing for file: {file.filename}, UID: {uid}, Source: {source}")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
        content = await file.read()
        import io
        Image.open(io.BytesIO(content)).convert('RGB').save(tmp.name, 'JPEG')
        path = tmp.name
    logger.info(f"Image file '{file.filename}' saved to temporary path: {path}")

    metadata = {"source": source, "id": str(uuid.uuid4()), "uid": uid}
    document = Document(content=path, metadata=metadata)
    
    logger.info(f"Storing image from source: '{source}' in ChromaDB.")
    db_image.store([document])
    logger.info(f"Successfully stored image from source: '{source}'.")
    os.remove(path)
    logger.info(f"Finished image processing and removed temporary file for: {file.filename}")

async def store_docx(file: UploadFile, uid: str, source: str):
    logger.info(f"Starting DOCX processing for file: {file.filename}, UID: {uid}, Source: {source}")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(await file.read())
        path = tmp.name
    logger.info(f"DOCX file '{file.filename}' saved to temporary path: {path}")

    doc = docx.Document(path)
    full_text = [para.text for para in doc.paragraphs]
    logger.info(f"Extracted text from '{file.filename}'")
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=20,
        length_function=len,
        is_separator_regex=False,
        separators=["\n\n", "\n", ".", ",", "\u200b", "\uff0c", "\u3001", "\uff0e", "\u3002", ""]
    )
    
    chunks = []
    split_docs = splitter.create_documents(["\n".join(full_text)])

    for j, chunk in enumerate(split_docs, start=1):
        metadata = {"source": source, "uid": uid, "id": str(uuid.uuid4()), "chunk": str(j)}
        document = Document(content=chunk.page_content, metadata=metadata)
        chunks.append(document)

    if chunks:
        logger.info(f"Storing {len(chunks)} chunks for source: '{source}' in ChromaDB.")
        db_docx.store(chunks)
        logger.info(f"Successfully stored chunks for source: '{source}'.")
    os.remove(path)
    logger.info(f"Finished DOCX processing and removed temporary file for: {file.filename}")

@router.post("/file/store")
async def store(uid: str = Form(...), source: str = Form(...), file: UploadFile = File(...)):
    logger.info(f"Received request to store 1 file for UID: {uid}, Source: {source}")
    
    content_type = file.content_type
    logger.info(f"Processing file: '{file.filename}', Content-Type: {content_type}")
    if content_type == "application/pdf":
        await store_pdf(file, uid, source)
    elif content_type.startswith("image/"):
        await store_image(file, uid, source)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        await store_docx(file, uid, source)
    else:
        logger.warning(f"File type '{content_type}' not supported for file '{file.filename}'.")
        raise HTTPException(status_code=400, detail=f"File type {content_type} not supported.")
            
    logger.info(f"Finished processing file for UID: {uid}, Source: {source}.")
    return {"answer": 'success'} 

@router.post("/query")
async def query(request: QueryRequest):
    logger.info(f"Received request to query for UID: {request.uid}, Query: {request.query}")
    
    results = db_pdf.query(request.query, k=5)
    return {"results": results}